from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

ROOT=r"D:\Project\Endless zombie"
OUT=os.path.join(ROOT,"output","report")
DIAG=os.path.join(OUT,"diagrams")
DOCX=os.path.join(OUT,"Project2_GD1305_Endless_Zombie_Report_EN.docx")
NAVY='1C3660'; BLUE='3063B2'; TEAL='2F92A8'; GOLD='E0A936'; LIGHT='EEF3F8'

def shade(cell, fill):
    pr=cell._tc.get_or_add_tcPr(); x=OxmlElement('w:shd'); x.set(qn('w:fill'),fill); pr.append(x)

def table(doc, headers, rows, fs=8.5):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,NAVY)
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(fs)
    trPr=t.rows[0]._tr.get_or_add_trPr(); rh=OxmlElement('w:tblHeader'); rh.set(qn('w:val'),'true'); trPr.append(rh)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=str(v); cs[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cs[i].paragraphs:
                for r in p.runs: r.font.size=Pt(fs)
    return t

def field(p, instruction):
    r=p.add_run(); a=OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'),'begin'); b=OxmlElement('w:instrText'); b.set(qn('xml:space'),'preserve'); b.text=instruction; c=OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'),'end'); r._r.extend([a,b,c])

def bullet(doc,text): doc.add_paragraph(text,style='List Bullet')
def step(doc,text): doc.add_paragraph(text,style='List Number')

def code_block(doc, path, language):
    name=os.path.basename(path)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(f'{name}  |  {language} source'); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
    text=open(path,encoding='utf-8').read().strip()
    p=doc.add_paragraph(); p.style=doc.styles['Normal']; p.paragraph_format.left_indent=Inches(.15); p.paragraph_format.right_indent=Inches(.15); p.paragraph_format.space_after=Pt(8)
    pp=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),'F4F6F8'); pp.append(sh)
    for i,line in enumerate(text.splitlines()):
        r=p.add_run(line + ('\n' if i<len(text.splitlines())-1 else '')); r.font.name='Consolas'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Consolas'); r.font.size=Pt(7.2); r.font.color.rgb=RGBColor(35,39,47)

def heading(doc,text,level=1): doc.add_heading(text,level=level)

def build():
    doc=Document(); s=doc.sections[0]; s.page_height=Inches(11.69); s.page_width=Inches(8.27); s.top_margin=Inches(.7); s.bottom_margin=Inches(.65); s.left_margin=Inches(.8); s.right_margin=Inches(.7)
    normal=doc.styles['Normal']; normal.font.name='Arial'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial'); normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.15
    for n,size,col in [('Title',28,NAVY),('Heading 1',17,NAVY),('Heading 2',13,BLUE),('Heading 3',11,TEAL)]:
        st=doc.styles[n]; st.font.name='Arial'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(col); st.paragraph_format.space_before=Pt(10); st.paragraph_format.space_after=Pt(6)
    h=s.header.paragraphs[0]; h.text='GD1305  |  PROJECT II  |  ENDLESS ZOMBIE'; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in h.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor(90,100,115)
    f=s.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.add_run('GD1305 - Endless Zombie     '); field(f,'PAGE')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(85); r=p.add_run('VTC ACADEMY'); r.bold=True; r.font.size=Pt(25); r.font.color.rgb=RGBColor.from_string(NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(55); r=p.add_run('PROJECT REPORT'); r.bold=True; r.font.size=Pt(34)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('ENDLESS ZOMBIE'); r.bold=True; r.font.size=Pt(32); r.font.color.rgb=RGBColor.from_string(NAVY)
    p=doc.add_paragraph('Survival Shooter / Action Roguelite'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    table(doc,['Information','Details'],[['Semester','Project II'],['Class','GD1305'],['Group','The MoonStone'],['Instructor','Nguyen Dinh Cuong'],['Group Member','Vu Viet Tien'],['Engine','Unity 6000.3.10f1']],fs=11)
    doc.add_paragraph('\nThis report follows the structure of the supplied Project II report form and is based on the current Endless Zombie source code, prefabs, and configuration assets.',style='Caption')
    doc.add_page_break(); heading(doc,'TABLE OF CONTENTS',0); p=doc.add_paragraph(); field(p,'TOC \\o "1-3" \\h \\z \\u'); doc.add_page_break()

    heading(doc,'I. PROJECT INTRODUCTION',0)
    heading(doc,'1. Project Overview',1)
    doc.add_paragraph('Endless Zombie is a top-down survival shooter in which the player controls an armed survivor inside an arena and attempts to survive increasingly difficult enemy waves. The combat loop combines movement, automatic target acquisition, weapon behavior, experience collection, level-up choices, and persistent upgrades purchased from the main menu.')
    doc.add_paragraph('The project uses a hybrid architecture. GameObject-oriented code handles input, menus, Canvas UI, audio, presentation, and game states. Unity Entities/DOTS handles wave scheduling, entity spawning, movement, projectiles, damage, combat events, and XP. Bridge systems synchronize Entity state with GameObject models, Animator controllers, and visual effects.')
    heading(doc,'2. Project Scope',1)
    for x in ['A Canvas-based main menu for stage selection, loadout selection, meta upgrades, and settings.','Player movement, automatic aiming, and several weapon profiles, including a multi-projectile Shotgun.','Data-driven stage and wave spawning through StageConfig and EnemyCatalog assets.','Regular zombie and DogMutant visual variants with movement-to-animation speed synchronization.','A Canvas-based gameplay HUD for health, XP, level, wave status, currency, and weapon statistics.','Running, paused, level-up, game-over, and win states.','Local persistence for music, sound, vibration, gold, and meta progression.']: bullet(doc,x)
    heading(doc,'3. System Name',1); doc.add_paragraph('Endless Zombie - Hybrid ECS Survival Shooter.')
    heading(doc,'4. Deployment Environment',1)
    table(doc,['Item','Configuration'],[['Target platform','Windows PC; responsive UI bootstrap is also available for mobile layouts'],['Engine','Unity 6000.3.10f1'],['Render pipeline','Universal Render Pipeline 17.3.0'],['Data-oriented stack','Unity Entities 1.4.2, Entities Graphics 1.4.15, Unity Physics 1.4.2'],['Language','C#'],['Input','Unity Input System 1.18.0']])
    heading(doc,'5. Development Tools',1); table(doc,['Tool','Purpose'],[['Unity Editor','Scenes, prefabs, animation, physics, profiling, and builds'],['Visual Studio / Rider','C# development and debugging'],['Git','Version control'],['TextMesh Pro / UGUI','Canvas menus and HUD'],['ScriptableObject assets','Stage, enemy, weapon, visual, and upgrade configuration']])
    heading(doc,'6. System Features',1); doc.add_paragraph('The implemented feature set includes data-driven waves, prefab-based enemy entities, chase behavior, projectiles, damage and status effects, XP and level-up upgrades, weapon profiles, persistent progression, editable Canvas UI, audio controls, and game completion states.')

    doc.add_page_break(); heading(doc,'II. SYSTEM REQUIREMENTS ANALYSIS',0)
    heading(doc,'1. System Overview',1)
    doc.add_paragraph('The system is organized into four responsibility groups: Presentation, OOP Gameplay, ECS Simulation, and Configuration/Persistence. This division keeps menu and HUD layout editable in the Unity Inspector while moving high-volume gameplay data into cache-friendly ECS systems.')
    heading(doc,'2. Functional Requirements',1)
    table(doc,['ID','Requirement','Priority'],[['FR-01','Start a selected stage from the main menu and load GameScene.','Must'],['FR-02','Move the player and automatically acquire a valid target.','Must'],['FR-03','Spawn enemies by wave, prefab, schedule, and maximum-alive limit.','Must'],['FR-04','Process projectiles, collisions, damage, death, and XP.','Must'],['FR-05','Present and apply level-up upgrade choices.','Must'],['FR-06','Display an Inspector-editable Canvas HUD.','Must'],['FR-07','Pause, resume, and configure music, sound, and vibration.','Should'],['FR-08','Persist gold and meta upgrades between sessions.','Should'],['FR-09','Finish a stage with either Win or Game Over.','Must']],fs=8.2)
    heading(doc,'3. Non-functional Requirements',1)
    for x in ['Stable performance with large groups of Entity enemies and limited unnecessary GameObject allocation.','UI layout must be authored as Canvas prefabs instead of being constructed by runtime layout code.','Gameplay values must be separated from logic through ScriptableObject configuration.','Visual run speed must follow Entity movement speed to reduce foot sliding.','Local settings must have safe defaults when saved values do not exist.','Wave configuration must validate missing IDs, duplicates, and invalid spawn entries.']: bullet(doc,x)
    heading(doc,'4. Use Case Diagram',1)
    doc.add_paragraph('The Use Case diagram is provided as PlantUML source so it can be regenerated or edited without redrawing it. The source is included in Appendix A and delivered as 01_use_case.puml.')
    heading(doc,'5. Use Case Descriptions',1)
    cases=[('UC-01','Start a Stage','Player','A valid stage and loadout are available.','Select stage and weapon, press Start, initialize the stage session, and load GameScene.','The player, HUD, and ECS stage runtime are active.'),('UC-02','Move and Auto Attack','Player','The game state is Running.','Read movement input, update the player, select a target in range, and fire according to the active weapon profile.','Projectiles are created and damage can be applied.'),('UC-03','Spawn an Enemy Wave','Game System','The stage has started or the previous wave is complete.','Activate a wave, schedule spawn entries, enqueue requests, and instantiate enemy Entity prefabs.','The configured enemies are active and counted.'),('UC-04','Level Up and Choose an Upgrade','Player','The player has collected enough XP.','Enter LevelUpState, show upgrade cards, apply the selected upgrade, and resume gameplay.','The chosen statistic or weapon modifier is updated.'),('UC-05','Pause and Configure Settings','Player','The application is running.','Open Settings, adjust music/sound/vibration, save values, and close or resume.','Settings are applied and stored in PlayerPrefs.'),('UC-06','Complete a Stage','Game System','The final wave is terminal and no enemies remain alive.','WaveCompletion confirms completion and emits StageCompleted.','The state machine enters WinState.')]
    for cid,name,actor,pre,flow,post in cases:
        heading(doc,f'{cid} - {name}',2); table(doc,['Field','Description'],[['Actor',actor],['Precondition',pre],['Main flow',flow],['Postcondition',post],['Exception','Missing or invalid configuration is logged and the invalid entry must not block stage completion permanently.']],fs=8.2)
    heading(doc,'6. Activity Diagram',1); doc.add_paragraph('The complete gameplay activity flow is supplied as Mermaid source in 02_gameplay_activity.mmd. It covers menu entry, stage initialization, wave activation, combat, wave repetition, Game Over, and Win branches.')

    doc.add_page_break(); heading(doc,'III. DETAILED DESIGN',0)
    heading(doc,'1. UI Design',1)
    table(doc,['Screen / Prefab','Main elements','Design responsibility'],[['MainMenuCanvas.prefab','Stage title, upgrade cards, weapon cards, settings, Start','MainMenuManager binds data and events; the prefab owns layout.'],['GameplayHUDLayout.prefab','HP, XP, level, wave, gold, weapon stats, settings','HUDManager updates values; the Canvas prefab owns layout.'],['SettingsMenu.prefab','Music slider, Sound slider, vibration, close','Fill and Handle use matching track padding and remain clamped to the slider bounds.'],['LevelUpMenu.prefab','Upgrade panels','Displays upgrade choices while LevelUpState is active.'],['Game Over / Win UI','Result and navigation actions','Activated by the game state machine.']],fs=8.2)
    doc.add_paragraph('The visual system uses a dark blue/purple base with yellow accents. TextMesh Pro provides readable typography. The Start button is visually transparent so only its label remains visible.')
    heading(doc,'2. Hybrid Architecture Diagram',1); doc.add_paragraph('The OOP/ECS boundary and bridge responsibilities are defined in Mermaid source 03_hybrid_architecture.mmd.')
    heading(doc,'3. Class / Component Diagram',1); doc.add_paragraph('The main classes, ECS components, and dependencies are defined in Mermaid class diagram source 04_class_component.mmd.')
    heading(doc,'4. Sequence Diagram',1); doc.add_paragraph('The spawn-to-combat sequence is defined in Mermaid sequence diagram source 05_spawn_combat_sequence.mmd.')
    heading(doc,'5. Gameplay Design',1)
    heading(doc,'5.1. Wave and Spawn Pipeline',2); doc.add_paragraph('StageConfig contains WaveDefinition records; each wave contains SpawnEntryDefinition records that reference EnemyId values in EnemyCatalog. WaveSpawnAuthoring bakes the authoring configuration into StageRuntime and dynamic buffers. SpawnEntrySchedulerSystem creates SpawnRequest items according to quantity and interval, while SpawnRequestProcessingSystem instantiates the resolved Entity prefab at the requested location. WaveCompletionSystem completes a wave only when all entries are terminal and the alive count reaches zero.')
    heading(doc,'5.2. Weapons and Shotgun Spread',2); doc.add_paragraph('WeaponManager stores damage, projectile count, spread angle, cooldown, range, and projectile speed. PlayerAutoAttackSystem distributes multiple shot directions across the configured cone. The Shotgun therefore fires a pellet group instead of a single projectile. GunStatsSystem combines base values with runtime modifiers from upgrades and meta progression.')
    heading(doc,'5.3. Enemy Visuals and Animation',2); doc.add_paragraph('Mob entities own simulation data such as health, damage, and movement speed. MobVisualBridge creates a matching GameObject model and synchronizes its transform every frame. MobVisualVariantAuthoring selects the DogMutant presentation for the dog enemy. Animator speed is computed from actual movement speed, run-loop duration, and distance travelled per animation loop, reducing visible foot sliding. Visual scale and ground offset remain configurable independently of the Entity collider.')
    heading(doc,'5.4. State Machine',2); doc.add_paragraph('GameStateMachineRunner coordinates Running, PlayerPaused, LevelUp, GameOver, and Win. StageCompleted moves the game to WinState; zero player health moves it to GameOverState; an XP threshold temporarily moves it to LevelUpState.')
    heading(doc,'6. Data Model',1); doc.add_paragraph('The project does not use a database server. ScriptableObject assets represent configuration relationships and PlayerPrefs stores local settings and progression. The model is supplied as Mermaid ER source in 06_data_model.mmd.')
    table(doc,['Entity','Representative attributes','Storage'],[['StageConfig','StageId, DefaultWaveDelay, MaxAliveEnemies, Waves','ScriptableObject'],['WaveDefinition','WaveId, WaveType, ActivationCondition, WaveDelay, Entries','Nested configuration'],['SpawnEntryDefinition','EnemyId, Quantity, SpawnInterval, spawn position/area','Nested configuration'],['EnemyCatalogEntry','EnemyId, EntityPrefab','EnemyCatalog.asset'],['GunConfig','Damage, cooldown, projectile count, spread, range','ScriptableObject'],['MobVisualSettings','Prefabs, controllers, loop timing, scale/offset data','Resources asset'],['PlayerPrefs','Music, sound, vibration, gold, meta upgrades','Local machine']],fs=8.1)

    doc.add_page_break(); heading(doc,'IV. TESTING',0)
    doc.add_paragraph('The table separates static project inspection from Unity Play Mode verification. “Pass - static” confirms that the required code, prefab, or configuration path exists. “Play Mode required” indicates that runtime behavior still needs final execution evidence on the submission machine.')
    tests=[('TC-01','Main Menu Canvas','Open MainMenu.','The prefab loads and buttons are bound.','Pass - static'),('TC-02','Gameplay HUD Canvas','Load GameScene.','HUD values are displayed and layout remains prefab-editable.','Pass - static'),('TC-03','Settings sliders','Drag Music and Sound to both ends.','Fill follows Handle; Handle stays inside track; volume follows value.','Pass - static'),('TC-04','Player movement','Use movement controls in Running.','The character moves smoothly in the requested direction.','Play Mode required'),('TC-05','Wave spawning','Begin a stage.','Correct prefabs, quantity, interval, and alive limit are used.','Pass - static'),('TC-06','Enemy ground contact','Observe spawned enemies.','Collider and visual offset place the model on the ground.','Play Mode required'),('TC-07','DogMutant animation','Spawn the dog variant.','Run animation plays and speed follows movement speed.','Pass - static'),('TC-08','Shotgun spread','Equip and fire the Shotgun.','Multiple projectiles form a configured spread cone.','Pass - static'),('TC-09','XP and level up','Collect enough XP.','Upgrade selection appears and modifies the chosen stat.','Pass - static'),('TC-10','Pause and resume','Open Settings/Pause and resume.','Simulation pauses and continues in the correct state.','Play Mode required'),('TC-11','Win condition','Clear the final wave.','StageCompleted is emitted and WinState is entered.','Pass - static'),('TC-12','Persistence','Change settings and restart.','PlayerPrefs restore the previous values.','Pass - static')]
    table(doc,['ID','Area','Action','Expected result','Status'],tests,fs=7.3)
    heading(doc,'Acceptance Criteria',1)
    for x in ['No compile errors in the Unity Console.','Enemy models do not flash at the player position during spawn.','DogMutant shows no obvious foot sliding at default and modified movement speeds.','Slider fill, handle, and volume value remain synchronized in the [0,1] range.','The final wave can complete even when an invalid spawn entry has been marked terminal.']: bullet(doc,x)

    doc.add_page_break(); heading(doc,'V. TASK ASSIGNMENT',0)
    table(doc,['No.','Task','Owner','Status'],[['1','Requirement analysis and architecture design','Vu Viet Tien','Completed'],['2','OOP/ECS gameplay, waves, combat, and XP','Vu Viet Tien','Completed for the current version'],['3','Main Menu, HUD, and Settings Canvas','Vu Viet Tien','Completed for the current version'],['4','Enemy and DogMutant visuals/animation','Vu Viet Tien','Completed for the current version'],['5','Testing, bug fixing, and report writing','Vu Viet Tien','Final Play Mode regression pending']],fs=8.8)
    heading(doc,'Self-assessment',1); doc.add_paragraph('The project demonstrates a practical hybrid Unity architecture, data-driven configuration, and editor-friendly Canvas workflows. The highest-priority action before submission is a complete Play Mode regression pass with screenshots and a Unity Profiler capture from a high-enemy-count wave.')

    doc.add_page_break(); heading(doc,'VI. INSTALLATION INSTRUCTIONS',0)
    heading(doc,'1. Prerequisites',1); doc.add_paragraph('Unity Hub, Unity Editor 6000.3.10f1, Windows 10/11, and optionally Git.')
    heading(doc,'2. Open the Project',1)
    for x in ['Open Unity Hub and choose Add project from disk.','Select the Endless zombie project directory.','Select Unity Editor 6000.3.10f1 and allow package restoration to complete.','Confirm that the Console contains no compile errors.']: step(doc,x)
    heading(doc,'3. Run in the Editor',1)
    for x in ['Open Assets/Scenes/MainMenu.unity.','Press Play.','Choose a stage and loadout, then press Tap to Start.','Verify the HUD, waves, Shotgun, DogMutant, pause, and settings behavior.']: step(doc,x)
    heading(doc,'4. Build for Windows',1); doc.add_paragraph('Open File > Build Profiles > Windows, add MainMenu and GameScene to the Scene List, and choose Build. Keep the generated executable and its associated data directory together.')
    heading(doc,'VII. CONCLUSION AND FUTURE DEVELOPMENT',0)
    doc.add_paragraph('Endless Zombie implements the complete survival-shooter loop: stage entry, wave spawning, automatic combat, XP and upgrades, and win/loss completion. The hybrid architecture provides efficient simulation while preserving a visual Unity workflow for UI, models, animation, and VFX.')
    for x in ['Add bosses and more enemy archetypes with explicit attack state machines.','Introduce pooling for visual bridges and VFX, then profile CPU time and garbage collection.','Expand stages, weapon synergies, and meta progression.','Add automated Play Mode tests for spawning, win conditions, and persistence.','Improve accessibility with input remapping, contrast options, text scaling, and reduced screen shake.']: bullet(doc,x)

    doc.add_page_break(); heading(doc,'APPENDIX A. DIAGRAM SOURCE CODE',0)
    doc.add_paragraph('Use Case is authored in PlantUML. Every other diagram is authored in Mermaid. The same source files are delivered beside this report in the diagrams folder and can be pasted directly into PlantUML or Mermaid-compatible editors.')
    code_block(doc,os.path.join(DIAG,'01_use_case.puml'),'PlantUML')
    doc.add_page_break(); code_block(doc,os.path.join(DIAG,'02_gameplay_activity.mmd'),'Mermaid flowchart'); code_block(doc,os.path.join(DIAG,'03_hybrid_architecture.mmd'),'Mermaid flowchart')
    doc.add_page_break(); code_block(doc,os.path.join(DIAG,'04_class_component.mmd'),'Mermaid classDiagram'); code_block(doc,os.path.join(DIAG,'05_spawn_combat_sequence.mmd'),'Mermaid sequenceDiagram')
    doc.add_page_break(); code_block(doc,os.path.join(DIAG,'06_data_model.mmd'),'Mermaid erDiagram')
    heading(doc,'REFERENCES',0)
    for x in ['Unity Technologies. Unity Manual 6000.x.','Unity Technologies. Entities 1.4 documentation.','Unity Technologies. Universal Render Pipeline 17 documentation.','Endless Zombie source code, prefabs, and ScriptableObject assets, reviewed August 2026.','Project2_GD1305_VuVietTien.pdf, used only as the report-structure reference.']: doc.add_paragraph(x)
    doc.core_properties.title='Project Report - Endless Zombie'; doc.core_properties.author='Vu Viet Tien'; doc.core_properties.subject='GD1305 Project II'; doc.core_properties.keywords='Unity, ECS, DOTS, Endless Zombie, PlantUML, Mermaid'
    doc.save(DOCX); print(DOCX)

if __name__=='__main__': build()
