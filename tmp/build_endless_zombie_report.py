from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import os, textwrap

ROOT = r"D:\Project\Endless zombie"
OUT = os.path.join(ROOT, "output", "report")
DIAG = os.path.join(ROOT, "tmp", "report_diagrams")
os.makedirs(OUT, exist_ok=True); os.makedirs(DIAG, exist_ok=True)
DOCX = os.path.join(OUT, "Project2_GD1305_Endless_Zombie_Report.docx")
NAVY=(28,54,96); BLUE=(48,99,178); CYAN=(47,146,168); GOLD=(224,169,54); PALE=(237,243,250); DARK=(35,39,47); GRAY=(95,105,118); RED=(180,65,65); GREEN=(54,130,92); WHITE=(255,255,255)

def font(size=28,bold=False):
    paths=[r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",r"C:\Windows\Fonts\calibri.ttf"]
    for p in paths:
        if os.path.exists(p): return ImageFont.truetype(p,size)
    return ImageFont.load_default()

def box(d, xy, text, fill=PALE, outline=BLUE, fs=25, radius=16, width=3):
    d.rounded_rectangle(xy,radius=radius,fill=fill,outline=outline,width=width)
    x1,y1,x2,y2=xy; f=font(fs,True); maxw=x2-x1-24
    words=text.split(); lines=[]; cur=''
    for w in words:
        test=(cur+' '+w).strip()
        if d.textbbox((0,0),test,font=f)[2] <= maxw: cur=test
        else: lines.append(cur); cur=w
    if cur: lines.append(cur)
    total=len(lines)*(fs+7); y=(y1+y2-total)/2
    for line in lines:
        bb=d.textbbox((0,0),line,font=f); d.text(((x1+x2-(bb[2]-bb[0]))/2,y),line,font=f,fill=DARK); y+=fs+7

def arrow(d, a, b, color=GRAY, width=4):
    d.line([a,b],fill=color,width=width)
    import math
    ang=math.atan2(b[1]-a[1],b[0]-a[0]); L=14
    p1=(b[0]-L*math.cos(ang-.55),b[1]-L*math.sin(ang-.55)); p2=(b[0]-L*math.cos(ang+.55),b[1]-L*math.sin(ang+.55))
    d.polygon([b,p1,p2],fill=color)

def title(d, text, sub=''):
    d.text((60,35),text,font=font(38,True),fill=NAVY)
    if sub: d.text((60,84),sub,font=font(20),fill=GRAY)
    d.line((60,122,1540,122),fill=GOLD,width=5)

def save(img,name):
    p=os.path.join(DIAG,name); img.save(p); return p

def diagrams():
    paths=[]
    im=Image.new('RGB',(1600,950),WHITE); d=ImageDraw.Draw(im); title(d,'USE CASE — ENDLESS ZOMBIE','Tác nhân chính: Người chơi; tác nhân phụ: Unity runtime')
    box(d,(70,330,300,550),'NGƯỜI CHƠI',fill=(245,245,245),outline=NAVY,fs=30)
    cases=[('Bắt đầu trận',(440,170,760,270)),('Di chuyển & ngắm bắn',(440,310,760,410)),('Tự động tấn công',(440,450,760,550)),('Chọn nâng cấp',(440,590,760,690)),('Tạm dừng / Cài đặt',(440,730,760,830)),('Sinh wave quái',(1000,170,1350,270)),('Nhận XP / lên cấp',(1000,370,1350,470)),('Thắng / thua trận',(1000,570,1350,670)),('Lưu meta progression',(1000,750,1350,850))]
    for t,xy in cases: box(d,xy,t,fill=PALE,outline=BLUE,fs=23)
    for y in [220,360,500,640,780]: arrow(d,(300,440),(440,y))
    arrow(d,(760,220),(1000,220)); arrow(d,(760,500),(1000,420)); arrow(d,(760,640),(1000,420)); arrow(d,(1175,470),(1175,570)); arrow(d,(1175,670),(1175,750))
    paths.append(save(im,'01_use_case.png'))

    im=Image.new('RGB',(1600,1000),WHITE); d=ImageDraw.Draw(im); title(d,'ACTIVITY DIAGRAM — VÒNG LẶP TRẬN ĐẤU')
    nodes=[('Main Menu',(600,145,1000,225)),('Nạp GameScene & dữ liệu Stage',(600,275,1000,355)),('Khởi tạo Player, HUD, ECS World',(600,405,1000,485)),('Kích hoạt Wave',(600,535,1000,615)),('Spawn → Chase → Attack',(600,665,1000,745)),('Wave đã hết quái?',(600,795,1000,875))]
    for i,(t,xy) in enumerate(nodes): box(d,xy,t,fill=(247,250,253),outline=BLUE,fs=23); 
    for i in range(len(nodes)-1): arrow(d,(800,nodes[i][1][3]),(800,nodes[i+1][1][1]))
    box(d,(1100,795,1450,875),'Còn wave?',fill=(255,249,232),outline=GOLD,fs=23); arrow(d,(1000,835),(1100,835)); arrow(d,(1275,795),(1000,575),color=GOLD)
    box(d,(1100,900,1450,975),'WIN',fill=(235,248,239),outline=GREEN,fs=26); arrow(d,(1275,875),(1275,900),color=GREEN)
    box(d,(120,665,470,745),'HP Player = 0?',fill=(255,239,239),outline=RED,fs=23); arrow(d,(600,705),(470,705),color=RED); box(d,(120,825,470,905),'GAME OVER',fill=(255,239,239),outline=RED,fs=26); arrow(d,(295,745),(295,825),color=RED)
    paths.append(save(im,'02_activity.png'))

    im=Image.new('RGB',(1600,980),WHITE); d=ImageDraw.Draw(im); title(d,'KIẾN TRÚC HỆ THỐNG HYBRID OOP + ECS/DOTS')
    d.rounded_rectangle((60,150,500,900),20,fill=(244,248,253),outline=BLUE,width=4); d.text((170,175),'OOP / GAMEOBJECT',font=font(28,True),fill=NAVY)
    for i,t in enumerate(['MainMenuManager','GameManager + State Machine','PlayerInput / CharacterLogic','HUD & Settings Canvas','AudioManager / MetaProgression']): box(d,(100,250+i*115,460,330+i*115),t,fill=WHITE,outline=BLUE,fs=20)
    d.rounded_rectangle((1100,150,1540,900),20,fill=(239,250,247),outline=GREEN,width=4); d.text((1210,175),'ECS / DOTS',font=font(28,True),fill=GREEN)
    for i,t in enumerate(['Wave Progression','Spawn Scheduler','Unit Mover / Chase','Projectile & Damage','XP / Combat Metrics']): box(d,(1140,250+i*115,1500,330+i*115),t,fill=WHITE,outline=GREEN,fs=20)
    d.rounded_rectangle((590,250,1010,800),20,fill=(255,250,237),outline=GOLD,width=4); d.text((700,275),'BRIDGE LAYER',font=font(28,True),fill=(145,102,20))
    for i,t in enumerate(['MobVisualBridge','DamageBridge','XPBridge','WeaponVfxBridge']): box(d,(630,360+i*105,970,435+i*105),t,fill=WHITE,outline=GOLD,fs=20)
    for y in [390,495,600,705]: arrow(d,(500,y),(590,y),color=GOLD); arrow(d,(1010,y),(1100,y),color=GOLD)
    paths.append(save(im,'03_architecture.png'))

    im=Image.new('RGB',(1600,1050),WHITE); d=ImageDraw.Draw(im); title(d,'CLASS / COMPONENT DIAGRAM','Các quan hệ chính được rút gọn để dễ đọc')
    classes=[('GameManager',['stateMachine','GameRunning / Paused / Over']),('HUDManager',['GameplayHUDView','Refresh HP / XP / Wave']),('MetaProgression',['StageUpgradeSnapshot','Gold & stage upgrades']),('WaveSpawnAuthoring',['StageConfig','EnemyCatalog','Baker → buffers']),('StageRuntime',['CurrentWaveIndex','State','SpawnRequests']),('Mob',['Health','MovementSpeed','Damage']),('WeaponManager',['ProjectileCount','SpreadAngle','AttackCooldown']),('MobVisualBridge',['VisualInstance','Animator speed sync'])]
    pos=[(70,170),(590,170),(1110,170),(70,500),(590,500),(1110,500),(330,790),(850,790)]
    for (name,attrs),(x,y) in zip(classes,pos):
        d.rounded_rectangle((x,y,x+410,y+200),14,fill=WHITE,outline=BLUE,width=3); d.rectangle((x,y,x+410,y+55),fill=NAVY); d.text((x+18,y+12),name,font=font(24,True),fill=WHITE)
        for j,a in enumerate(attrs): d.text((x+20,y+75+j*40),'• '+a,font=font(19),fill=DARK)
    arrow(d,(275,370),(275,500)); arrow(d,(795,370),(795,500)); arrow(d,(1315,370),(1315,500)); arrow(d,(480,890),(590,600)); arrow(d,(1055,790),(1315,700)); arrow(d,(850,890),(740,600))
    paths.append(save(im,'04_class_component.png'))

    im=Image.new('RGB',(1800,1050),WHITE); d=ImageDraw.Draw(im); title(d,'SEQUENCE DIAGRAM — SPAWN, TẤN CÔNG VÀ TIÊU DIỆT QUÁI')
    actors=['GameState','WaveProgression','SpawnScheduler','SpawnProcessor','Mob Systems','Bridge / HUD']
    xs=[150,450,750,1050,1350,1650]
    for x,a in zip(xs,actors): box(d,(x-115,145,x+115,215),a,fill=PALE,outline=BLUE,fs=18); d.line((x,215,x,990),fill=(185,190,198),width=2)
    msgs=[(0,1,'BeginStage()',270),(1,2,'Activate wave',360),(2,3,'Enqueue SpawnRequest',450),(3,4,'Instantiate enemy prefab',540),(4,5,'Create visual model',630),(4,4,'Chase + auto attack + damage',720),(4,5,'MobDeath / XP event',810),(5,1,'Update alive count',900),(1,0,'StageCompleted',970)]
    for a,b,t,y in msgs:
        if a==b: d.arc((xs[a],y-20,xs[a]+90,y+30),270,90,fill=GRAY,width=3); d.text((xs[a]+100,y-14),t,font=font(18),fill=DARK)
        else: arrow(d,(xs[a],y),(xs[b],y),color=GREEN if b>a else GOLD,width=3); mid=(xs[a]+xs[b])//2; bb=d.textbbox((0,0),t,font=font(18)); d.rectangle((mid-(bb[2]-bb[0])/2-5,y-27,mid+(bb[2]-bb[0])/2+5,y-3),fill=WHITE); d.text((mid-(bb[2]-bb[0])/2,y-27),t,font=font(18),fill=DARK)
    paths.append(save(im,'05_sequence.png'))

    im=Image.new('RGB',(1600,900),WHITE); d=ImageDraw.Draw(im); title(d,'DATA MODEL — CẤU HÌNH VÀ LƯU TRỮ CỤC BỘ')
    box(d,(80,180,430,300),'StageConfig.asset',fill=PALE,outline=BLUE,fs=26); box(d,(625,180,975,300),'WaveDefinition[]',fill=PALE,outline=BLUE,fs=26); box(d,(1170,180,1520,300),'SpawnEntryDefinition[]',fill=PALE,outline=BLUE,fs=23)
    arrow(d,(430,240),(625,240)); arrow(d,(975,240),(1170,240))
    box(d,(1170,400,1520,520),'EnemyCatalog.asset',fill=(239,250,247),outline=GREEN,fs=25); arrow(d,(1345,300),(1345,400),color=GREEN)
    box(d,(80,600,430,730),'PlayerPrefs',fill=(255,249,232),outline=GOLD,fs=27); box(d,(625,570,975,760),'GoldWallet\nMusic/Sound\nVibration',fill=(255,249,232),outline=GOLD,fs=24); box(d,(1170,600,1520,730),'MetaProgression',fill=(255,249,232),outline=GOLD,fs=25)
    arrow(d,(430,665),(625,665),color=GOLD); arrow(d,(975,665),(1170,665),color=GOLD)
    paths.append(save(im,'06_data_model.png'))
    return paths

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:tblHeader'); e.set(qn('w:val'),'true'); trPr.append(e)

def add_table(doc, headers, rows, widths=None, fs=9):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; set_cell_shading(c,'1C3660');
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(fs)
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.size=Pt(fs)
    return t

def add_field(paragraph, instr):
    run=paragraph.add_run(); fldChar=OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'),'begin'); instrText=OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text=instr; fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end'); run._r.extend([fldChar,instrText,fldChar2])

def make_doc(paths):
    doc=Document(); sec=doc.sections[0]; sec.page_height=Inches(11.69); sec.page_width=Inches(8.27); sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.8); sec.right_margin=Inches(.7)
    styles=doc.styles; styles['Normal'].font.name='Arial'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial'); styles['Normal'].font.size=Pt(10.5); styles['Normal'].paragraph_format.space_after=Pt(5); styles['Normal'].paragraph_format.line_spacing=1.15
    for n,size,col in [('Title',28,'1C3660'),('Heading 1',17,'1C3660'),('Heading 2',13,'3063B2'),('Heading 3',11,'2F92A8')]:
        s=styles[n]; s.font.name='Arial'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(10); s.paragraph_format.space_after=Pt(6)
    # header/footer
    h=sec.header.paragraphs[0]; h.text='GD1305  |  PROJECT II  |  ENDLESS ZOMBIE'; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in h.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor(90,100,115)
    f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.add_run('GD1305 – Endless Zombie     '); add_field(f,'PAGE')
    # cover
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(85); r=p.add_run('VTC ACADEMY'); r.bold=True; r.font.size=Pt(25); r.font.color.rgb=RGBColor(*NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(55); r=p.add_run('PROJECT REPORT'); r.bold=True; r.font.size=Pt(34)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('ENDLESS ZOMBIE'); r.bold=True; r.font.size=Pt(32); r.font.color.rgb=RGBColor(*NAVY)
    doc.add_paragraph('Survival Shooter / Action Roguelite',style='Subtitle').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    add_table(doc,['Thông tin','Nội dung'],[['Semester','Project II'],['Class','GD1305'],['Group','The MoonStone'],['Instructor','Nguyễn Đình Cường'],['Group Member','Vũ Việt Tiến'],['Engine','Unity 6000.3.10f1']],fs=11)
    doc.add_paragraph('\nTài liệu được xây dựng dựa trên cấu trúc form Project Report được cung cấp và đối chiếu với mã nguồn/prefab hiện có của dự án.',style='Caption')
    doc.add_page_break()
    doc.add_heading('MỤC LỤC',0); p=doc.add_paragraph(); add_field(p,'TOC \\o "1-3" \\h \\z \\u'); doc.add_paragraph('Trong Microsoft Word: nhấp phải vào mục lục → Update Field → Update entire table.',style='Caption'); doc.add_page_break()
    doc.add_heading('I. GIỚI THIỆU DỰ ÁN',0)
    doc.add_heading('1. Tổng quan dự án',1)
    doc.add_paragraph('Endless Zombie là trò chơi bắn súng sinh tồn góc nhìn từ trên xuống, trong đó người chơi điều khiển một chiến binh giữa bản đồ đấu trường, tự động ngắm/bắn mục tiêu gần nhất và sống sót qua các wave zombie có độ khó tăng dần. Mỗi quái bị tiêu diệt tạo ra XP; khi lên cấp, người chơi chọn nâng cấp để tạo build chiến đấu khác nhau. Vòng lặp chính kết hợp khả năng né tránh, kiểm soát vị trí, lựa chọn vũ khí và nâng cấp dài hạn ở Main Menu.')
    doc.add_paragraph('Điểm kỹ thuật nổi bật là kiến trúc hybrid: GameObject/OOP phụ trách trình bày, input, UI, âm thanh và state machine; Unity Entities/DOTS xử lý dữ liệu số lượng lớn như spawn wave, di chuyển, đạn, sát thương và sự kiện chiến đấu. Bridge systems đồng bộ Entity với model/Animator/VFX GameObject.')
    doc.add_heading('2. Phạm vi dự án',1)
    for x in ['Main Menu dạng Canvas: chọn stage, nâng cấp Max HP/Income, chọn vũ khí, mở Settings.','Gameplay: di chuyển, tự động tấn công, nhiều loại súng; Shotgun bắn đạn chùm theo projectile count và spread angle.','Wave system theo StageConfig/EnemyCatalog, giới hạn số quái sống và spawn từ prefab Entity.','Quái thường và DogMutant có visual riêng; animation tốc độ chạy được đồng bộ theo tốc độ di chuyển thực tế.','HUD dạng Canvas hiển thị HP, XP, level, wave, vàng và trạng thái vũ khí.','Game state: Running, PlayerPaused, LevelUp, GameOver và Win.','Lưu cục bộ âm lượng, rung, vàng và meta progression bằng PlayerPrefs.']:
        doc.add_paragraph(x,style='List Bullet')
    doc.add_heading('3. Tên hệ thống',1); doc.add_paragraph('Endless Zombie — Hybrid ECS Survival Shooter.')
    doc.add_heading('4. Môi trường triển khai',1); add_table(doc,['Hạng mục','Cấu hình'],[['Nền tảng mục tiêu','Windows PC; UI có bootstrap responsive cho thiết bị di động'],['Engine','Unity 6000.3.10f1'],['Render pipeline','Universal Render Pipeline 17.3.0'],['Kiến trúc dữ liệu','Unity Entities 1.4.2, Entities Graphics 1.4.15, Unity Physics 1.4.2'],['Ngôn ngữ','C#'],['Điều khiển','Unity Input System 1.18.0']],fs=9)
    doc.add_heading('5. Công cụ phát triển',1); add_table(doc,['Công cụ','Mục đích'],[['Unity Editor','Scene, prefab, animation, physics, build'],['Visual Studio / Rider','Lập trình và debug C#'],['Git','Quản lý phiên bản'],['TextMesh Pro / UGUI','Hiển thị HUD và menu Canvas'],['ScriptableObject','Cấu hình stage, wave, enemy, weapon và upgrade']],fs=9)
    doc.add_heading('6. Tính năng hệ thống',1); doc.add_paragraph('Các tính năng cốt lõi gồm spawn wave theo dữ liệu, giới hạn alive, AI đuổi mục tiêu, sát thương/hiệu ứng, XP và level-up, nhiều weapon profile, meta progression, Canvas UI và điều khiển âm lượng/rung.')
    doc.add_page_break()
    doc.add_heading('II. PHÂN TÍCH YÊU CẦU HỆ THỐNG',0)
    doc.add_heading('1. Tổng quan hệ thống',1); doc.add_paragraph('Hệ thống được chia thành bốn lớp: Presentation (Canvas, model, Animator, VFX, audio), Gameplay OOP (input, state machine, player services), Simulation ECS (wave, mob, projectile, damage, XP) và Configuration/Persistence (ScriptableObject, PlayerPrefs). Cách chia này giữ UI dễ chỉnh trong Inspector nhưng vẫn tận dụng xử lý data-oriented cho số lượng quái lớn.')
    doc.add_picture(paths[2],width=Inches(6.7)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Hình 1. Kiến trúc hybrid OOP + ECS/DOTS của Endless Zombie.',style='Caption').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('2. Yêu cầu chức năng',1)
    add_table(doc,['ID','Yêu cầu','Mức ưu tiên'],[['FR-01','Bắt đầu một stage từ Main Menu và nạp GameScene.','Must'],['FR-02','Điều khiển nhân vật và tự động chọn mục tiêu trong tầm.','Must'],['FR-03','Sinh quái theo wave, prefab, lịch spawn và giới hạn alive.','Must'],['FR-04','Tính đạn, va chạm, sát thương, chết và XP.','Must'],['FR-05','Cho phép chọn nâng cấp khi lên level.','Must'],['FR-06','Hiển thị HUD bằng Canvas prefab chỉnh được trong Inspector.','Must'],['FR-07','Tạm dừng, tiếp tục, chỉnh Music/Sound/Vibration.','Should'],['FR-08','Lưu vàng và nâng cấp meta giữa các phiên.','Should'],['FR-09','Kết thúc bằng Win hoặc Game Over.','Must']],fs=8.5)
    doc.add_heading('3. Yêu cầu phi chức năng',1)
    add_table(doc,['ID','Yêu cầu'],[['NFR-01','Hiệu năng ổn định khi có nhiều Entity; tránh tạo/hủy GameObject không cần thiết trong simulation.'],['NFR-02','UI responsive và không dựng layout bằng code; các thay đổi bố cục thực hiện trên Canvas prefab.'],['NFR-03','Cấu hình gameplay tách khỏi code bằng ScriptableObject.'],['NFR-04','Animation visual đồng bộ với vận tốc Entity, hạn chế trượt chân.'],['NFR-05','Dữ liệu cục bộ có giá trị mặc định an toàn khi chưa tồn tại.'],['NFR-06','Các hệ thống wave có validation cho ID trùng, entry sai và config thiếu.']],fs=8.5)
    doc.add_page_break(); doc.add_heading('4. Use Case Diagram',1); doc.add_picture(paths[0],width=Inches(6.75)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph('Hình 2. Use case tổng quát.',style='Caption').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('5. Mô tả Use Case',1)
    ucs=[('UC-01','Bắt đầu trận','Người chơi','Chọn Tap to Start/Play','Stage session được khởi tạo; GameScene được nạp.','Chọn stage → kiểm tra loadout → bấm Start → BeginStageSession → LoadScene.'),('UC-02','Di chuyển và chiến đấu','Người chơi','Game đang ở Running','Player cập nhật vị trí; mục tiêu trong tầm nhận đạn/sát thương.','Nhập hướng → CharacterLogic di chuyển → hệ thống tìm mục tiêu → WeaponManager phát đạn.'),('UC-03','Sinh wave quái','Game System','Stage bắt đầu hoặc wave trước hoàn tất','SpawnRequest được xử lý; prefab quái xuất hiện đúng vị trí.','WaveProgression kích hoạt → Scheduler xếp yêu cầu → Processor instantiate prefab → Mob systems chạy.'),('UC-04','Lên cấp và chọn nâng cấp','Người chơi','Đủ XP','Chỉ số/vũ khí được thay đổi; game quay lại Running.','Thu XP → CharacterXPManager tăng level → LevelUpState → chọn card → áp dụng upgrade.'),('UC-05','Tạm dừng và cài đặt','Người chơi','Đang chơi hoặc ở Main Menu','Âm lượng/rung được lưu; game tiếp tục hoặc giữ pause.','Mở Settings → kéo Music/Sound hoặc đổi Vibration → AudioManager cập nhật PlayerPrefs.'),('UC-06','Hoàn thành stage','Game System','Wave cuối đã hoàn tất và không còn quái sống','Chuyển WinState và hiển thị tiến trình.','WaveCompletion xác nhận terminal → phát StageCompleted → state machine chuyển Win.')]
    for uid,name,actor,pre,post,flow in ucs:
        doc.add_heading(f'{uid} – {name}',2); add_table(doc,['Thuộc tính','Mô tả'],[['Actor',actor],['Tiền điều kiện',pre],['Hậu điều kiện',post],['Luồng chính',flow],['Ngoại lệ','Thiếu prefab/config: ghi lỗi validation và không để entry lỗi khóa vĩnh viễn stage.']],fs=8.5)
    doc.add_page_break(); doc.add_heading('6. Activity Diagram',1); doc.add_picture(paths[1],width=Inches(6.7)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph('Hình 3. Activity diagram vòng lặp trận đấu.',style='Caption').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break(); doc.add_heading('III. THIẾT KẾ CHI TIẾT',0)
    doc.add_heading('1. Thiết kế UI',1)
    add_table(doc,['Màn hình / Prefab','Thành phần','Nguyên tắc'],[['MainMenuCanvas.prefab','Stage title, upgrade cards, weapon cards, settings, Start','Layout chỉnh trực tiếp trên Canvas; MainMenuManager chỉ bind sự kiện và dữ liệu.'],['GameplayHUDLayout.prefab','HP, XP, level, wave, gold, weapon stats, settings','HUDManager cập nhật text/value; không dựng UI bằng code.'],['SettingsMenu.prefab','Music slider, Sound slider, vibration, close','Fill Area và Handle Slide Area dùng cùng padding; fill bám handle và handle bị clamp trong track.'],['LevelUpMenu.prefab','Các UpgradePanel','Hiển thị lựa chọn nâng cấp khi LevelUpState hoạt động.'],['Game Over / Win','Kết quả và nút retry/menu','Được kích hoạt theo state machine.']],fs=8.5)
    doc.add_paragraph('Màu chủ đạo xanh tím/đen làm nền, vàng dùng cho điểm nhấn và thanh giá trị. TextMesh Pro đảm bảo chữ rõ; button Start có nền trong suốt để chỉ giữ phần chữ theo yêu cầu thiết kế.')
    doc.add_heading('2. Class / Component Diagram',1); doc.add_picture(paths[3],width=Inches(6.7)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph('Hình 4. Class/Component diagram rút gọn.',style='Caption').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break(); doc.add_heading('3. Sequence Diagram',1); doc.add_picture(paths[4],width=Inches(6.75)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph('Hình 5. Chuỗi xử lý spawn–combat–hoàn thành stage.',style='Caption').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('4. Thiết kế gameplay',1)
    doc.add_heading('4.1. Wave và spawn',2); doc.add_paragraph('StageConfig chứa danh sách WaveDefinition; mỗi wave chứa SpawnEntryDefinition tham chiếu EnemyId trong EnemyCatalog. Khi bake, WaveSpawnAuthoring chuyển cấu hình thành StageRuntime và các dynamic buffer. SpawnEntrySchedulerSystem tạo SpawnRequest theo interval/quantity; SpawnRequestProcessingSystem instantiate Entity prefab và áp vị trí spawn. WaveCompletionSystem chỉ hoàn tất wave khi toàn bộ entry terminal và số quái sống bằng 0.')
    doc.add_heading('4.2. Weapon và Shotgun',2); doc.add_paragraph('WeaponManager giữ projectile count, spread angle, cooldown, damage và projectile speed. PlayerAutoAttackSystem tạo nhiều hướng bắn phân bố quanh hướng mục tiêu. Với Shotgun, BaseProjectileCount lớn hơn 1 và BaseSpreadAngle tạo chùm đạn; các modifier nâng cấp được GunStatsSystem tổng hợp vào giá trị runtime.')
    doc.add_heading('4.3. Quái và visual bridge',2); doc.add_paragraph('Entity Mob lưu health, damage và movement speed; hệ thống chase/mover cập nhật LocalTransform. MobVisualBridge tạo model GameObject tương ứng và đồng bộ transform mỗi frame. DogMutant được chọn bằng MobVisualVariantAuthoring; tốc độ Animator được tính từ vận tốc thực tế, độ dài vòng chạy và quãng đường model đi trong một loop, giúp giảm hiện tượng trượt chân. GroundOffset và scale nằm trong MobVisualSettings/prefab visual, tách khỏi collider Entity.')
    doc.add_heading('4.4. Game state',2); doc.add_paragraph('GameStateMachineRunner điều phối Running, PlayerPaused, LevelUp, GameOver và Win. Sự kiện StageCompleted từ WaveSpawnLifecycle đưa hệ thống sang WinState; HP player bằng 0 đưa sang GameOverState; đủ XP chuyển tạm thời sang LevelUpState.')
    doc.add_heading('5. Thiết kế dữ liệu',1); doc.add_picture(paths[5],width=Inches(6.7)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph('Hình 6. Data model cấu hình và persistence.',style='Caption').alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc,['Thực thể','Thuộc tính tiêu biểu','Lưu trữ'],[['StageConfig','StageId, DefaultWaveDelay, MaxAliveEnemies, Waves[]','ScriptableObject asset'],['WaveDefinition','WaveId, WaveType, ActivationCondition, WaveDelay, Entries[]','Lồng trong StageConfig'],['SpawnEntryDefinition','EnemyId, Quantity, SpawnInterval, Position/Area','Lồng trong WaveDefinition'],['EnemyCatalogEntry','EnemyId, EntityPrefab','EnemyCatalog.asset'],['GunConfig','Damage, cooldown, projectile count, spread, range','ScriptableObject asset'],['MobVisualSettings','Visual prefab, controller, loop duration, distance/loop, ground offset','Resources asset'],['PlayerPrefs','Music, Sound, Vibration, Gold, meta upgrades','Local machine']],fs=8.5)
    doc.add_paragraph('Dự án không dùng hệ quản trị cơ sở dữ liệu. Quan hệ dữ liệu mang tính cấu hình nội bộ và lưu cục bộ; vì vậy sơ đồ trên thay thế ERD truyền thống nhưng vẫn thể hiện khóa tham chiếu EnemyId và quan hệ Stage → Wave → SpawnEntry.')
    doc.add_page_break(); doc.add_heading('IV. KIỂM THỬ',0)
    doc.add_paragraph('Bảng dưới phân biệt kiểm tra tĩnh (đối chiếu code/prefab/config) với play-test. Trạng thái “Đạt – tĩnh” nghĩa là cấu trúc triển khai đã được xác nhận trong project; cần chạy Unity Play Mode để chốt “Đạt – runtime” trên máy nộp bài.')
    tests=[('TC-01','Main Menu Canvas','Mở MainMenu','Prefab được load, button bind được sự kiện','Đạt – tĩnh'),('TC-02','Gameplay HUD Canvas','Nạp GameScene','HUD hiển thị HP/XP/wave và chỉnh bằng prefab','Đạt – tĩnh'),('TC-03','Settings slider','Kéo Music/Sound tới hai đầu','Fill bám handle; handle không ra ngoài track; volume thay đổi','Đạt – tĩnh'),('TC-04','Player movement','WASD/joystick trong Running','Nhân vật di chuyển đúng hướng','Cần Play Mode'),('TC-05','Wave spawning','BeginStage','Spawn đúng prefab, quantity, interval và max alive','Đạt – tĩnh'),('TC-06','Mob collider/ground','Quan sát quái spawn','Collider tồn tại; model chạm đất theo offset','Cần Play Mode'),('TC-07','DogMutant animation','Spawn ZombieDog variant','Model DogMutant chạy và Animator speed bám movement speed','Đạt – tĩnh'),('TC-08','Shotgun spread','Trang bị Shotgun và bắn','Tạo nhiều projectile dạng chùm theo spread angle','Đạt – tĩnh'),('TC-09','XP & Level Up','Diệt quái đến đủ XP','Thu XP, mở lựa chọn upgrade, áp chỉ số','Đạt – tĩnh'),('TC-10','Pause/Resume','Nhấn Esc/settings','Simulation dừng và tiếp tục đúng state','Cần Play Mode'),('TC-11','Win condition','Hoàn tất wave cuối','Phát StageCompleted và chuyển Win','Đạt – tĩnh'),('TC-12','Persistence','Đổi âm lượng/vàng rồi mở lại','PlayerPrefs khôi phục giá trị','Đạt – tĩnh')]
    add_table(doc,['ID','Hạng mục','Bước chính','Kết quả mong đợi','Trạng thái'],tests,fs=7.5)
    doc.add_heading('Tiêu chí nghiệm thu',1)
    for x in ['Không có lỗi compile trong Console.','Không xuất hiện model quái nhấp nháy tại vị trí player khi spawn.','DogMutant không trượt chân rõ rệt ở tốc độ mặc định và sau modifier.','Slider fill, handle và giá trị volume luôn đồng bộ trong [0,1].','Wave cuối kết thúc được ngay cả khi có entry cấu hình lỗi đã được đánh dấu terminal.']:
        doc.add_paragraph(x,style='List Bullet')
    doc.add_page_break(); doc.add_heading('V. PHÂN CÔNG CÔNG VIỆC',0)
    add_table(doc,['STT','Công việc','Phụ trách','Kết quả'],[['1','Phân tích yêu cầu và thiết kế kiến trúc','Vũ Việt Tiến','Hoàn thành'],['2','Gameplay OOP/ECS, wave, combat, XP','Vũ Việt Tiến','Hoàn thành theo phiên bản hiện tại'],['3','Main Menu, HUD, Settings Canvas','Vũ Việt Tiến','Hoàn thành theo phiên bản hiện tại'],['4','Quái, DogMutant visual/animation','Vũ Việt Tiến','Hoàn thành theo phiên bản hiện tại'],['5','Kiểm thử, sửa lỗi và viết báo cáo','Vũ Việt Tiến','Đang hoàn thiện play-test cuối']],fs=9)
    doc.add_heading('Đánh giá',1); doc.add_paragraph('Dự án thể hiện khả năng kết hợp quy trình Unity truyền thống với DOTS, tách cấu hình khỏi logic và chuyển UI sang prefab Canvas để dễ bảo trì. Phần cần ưu tiên trước khi nộp là chạy regression test đầy đủ trong Play Mode, ghi lại ảnh chụp UI thực tế và profile hiệu năng ở wave đông quái.')
    doc.add_page_break(); doc.add_heading('VI. HƯỚNG DẪN CÀI ĐẶT VÀ CHẠY',0)
    doc.add_heading('1. Yêu cầu',1); doc.add_paragraph('Unity Hub; Unity Editor 6000.3.10f1; Windows 10/11; Git (không bắt buộc nếu dùng file ZIP).')
    doc.add_heading('2. Mở project',1)
    for x in ['Mở Unity Hub → Add project from disk.','Chọn thư mục “Endless zombie”.','Đảm bảo Editor 6000.3.10f1 được chọn; Unity sẽ khôi phục package từ manifest.','Chờ hoàn tất import và không còn lỗi compile trong Console.']:
        doc.add_paragraph(x,style='List Number')
    doc.add_heading('3. Chạy trong Editor',1)
    for x in ['Mở Assets/Scenes/MainMenu.unity.','Nhấn Play.','Chọn stage/loadout, nhấn Tap to Start.','Kiểm tra HUD, wave, Shotgun, DogMutant, pause và settings.']:
        doc.add_paragraph(x,style='List Number')
    doc.add_heading('4. Build Windows',1); doc.add_paragraph('File → Build Profiles → Windows → bổ sung MainMenu và GameScene vào Scene List → Build. Chạy file .exe trong thư mục build; không tách thư mục *_Data khỏi file thực thi.')
    doc.add_heading('VII. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN',0)
    doc.add_paragraph('Endless Zombie đã hình thành đầy đủ vòng lặp survival shooter: bắt đầu stage, spawn wave, tự động chiến đấu, nhận XP/nâng cấp và kết thúc thắng/thua. Kiến trúc hybrid phù hợp với mục tiêu vừa có hiệu năng simulation vừa duy trì quy trình làm UI/model trực quan trong Unity Editor.')
    for x in ['Bổ sung boss và nhiều archetype quái có attack animation/state rõ ràng.','Thêm object pooling cho VFX/model bridge và đo CPU/GC bằng Unity Profiler.','Mở rộng stage, weapon synergy và meta progression.','Tạo automated Play Mode tests cho spawn, win condition và persistence.','Thêm accessibility: remap input, độ tương phản, cỡ chữ và tùy chọn giảm rung màn hình.']:
        doc.add_paragraph(x,style='List Bullet')
    doc.add_heading('TÀI LIỆU THAM CHIẾU',0)
    for x in ['Unity Technologies. Unity Manual 6000.x.','Unity Technologies. Entities package 1.4 documentation.','Unity Technologies. Universal Render Pipeline 17 documentation.','Mã nguồn, prefab và ScriptableObject trong project Endless Zombie (đối chiếu ngày 23/08/2026).','Project2_GD1305_VuVietTien.pdf — tài liệu tham chiếu bố cục báo cáo.']:
        doc.add_paragraph(x)
    doc.core_properties.title='Project Report – Endless Zombie'; doc.core_properties.author='Vũ Việt Tiến'; doc.core_properties.subject='GD1305 Project II'; doc.core_properties.keywords='Unity, ECS, DOTS, Endless Zombie, Project Report'
    doc.save(DOCX); return DOCX

if __name__=='__main__':
    print(make_doc(diagrams()))
